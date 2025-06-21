from typing import List
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from common.logger import setup_logger
from docx import Document as DocxDocument

logger = setup_logger('document_loader')

class DocxLoader:
    """python-docx를 사용한 DOCX 파일 로더"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """DOCX 파일을 로드하여 Document 리스트 반환"""
        try:
            doc = DocxDocument(self.file_path)
            text = ""
            
            # 모든 단락에서 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 표에서 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return [Document(page_content=text, metadata={"source": self.file_path})]
            
        except Exception as e:
            logger.error(f"DOCX 파일 로드 중 오류: {str(e)}")
            return []

def load_documents(file_path: str) -> List[Document]:
    """
    [문서 로딩 및 청크 분할]
    - 입력: 
        - file_path: 로드할 파일 경로
    - 출력: 분할된 문서 청크 리스트
    """
    try:
        # 파일 확장자에 따라 적절한 로더 선택
        if file_path.endswith('.txt'):
            loader = TextLoader(file_path, encoding='utf-8')
        elif file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
        elif file_path.endswith('.docx'):
            loader = DocxLoader(file_path)
        else:
            logger.error(f"지원하지 않는 파일 형식: {file_path}")
            return []

        # 문서 로드
        documents = loader.load()
        
        # 문서 분할
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            is_separator_regex=False
        )
        
        # 청크로 분할
        chunks = text_splitter.split_documents(documents)
        logger.info(f"문서 '{file_path}' 로드 완료: {len(chunks)}개 청크 생성")
        
        return chunks
        
    except Exception as e:
        logger.error(f"문서 로딩 중 오류 발생: {str(e)}")
        return [] 